from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\1\Desktop\家里用的图标")
DOCX_OUT = Path(r"C:\Users\1\Desktop\桌面文档\美国市场语言执行补充方案_20260322.docx")
MD_OUT = ROOT / "output" / "doc" / "美国市场语言执行补充方案_20260322.md"


def set_font(run, name="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def init_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    if "Small" not in doc.styles:
        style = doc.styles.add_style("Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(9)
    return doc


def add_title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, size=20, bold=True, color=(31, 78, 121))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11.5)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=14 if level == 1 else 12, bold=True, color=(31, 78, 121))


def add_para(doc, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    set_font(r, size=9 if style == "Small" else 10.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.35
        r1 = p.add_run("• ")
        set_font(r1)
        r2 = p.add_run(item)
        set_font(r2)


SUMMARY = [
    "这个生意可以做，语言不是前置门槛，关键是把“英语能力”从个人能力问题，改造成“流程 + 工具 + 模板 + 少量双语支持”的组织能力。",
    "老板不懂英文，不影响拍板；邀约岗位英文弱，也不等于不能做。真正不能接受的，是让他们直接裸奔去打英文电话、临场自由发挥。",
    "最稳的做法不是“全员学英语后再做”，而是先搭一个中文主导、英文输出的工作流：外部看起来是英文专业团队，内部决策和推进仍然用中文完成。",
]

STACK = [
    "网页阅读：Microsoft Edge 自带网页翻译。微软官方说明 Edge 会自动检测外语网页并提供翻译，支持 70 多种语言。[Microsoft Edge 翻译](https://support.microsoft.com/en-us/topic/use-microsoft-translator-in-microsoft-edge-browser-4ad1c6cb-01a4-4227-be9d-a81e127fcb0b)",
    "邮件和文档：DeepL + Microsoft 365。DeepL 官方说明支持 Word、Outlook、PowerPoint 插件，支持文档翻译和 glossary；Microsoft 官方说明 Word/Outlook 本身也支持翻译。[DeepL Word](https://support.deepl.com/hc/en-us/articles/9851524013468-DeepL-for-Microsoft-Word) [DeepL 文档翻译](https://www.deepl.com/en/translator/features/document-translation) [Outlook 翻译](https://support.microsoft.com/en-us/office/translator-in-outlook-for-windows-3d7e12ed-99d6-406e-a453-b9db0d9653fa)",
    "临时面对面或手机沟通：Microsoft Translator。微软官方现在仍提供多设备 conversation、split-screen 和浏览器加入会话能力，适合临时演示和现场沟通。[Microsoft Translator 功能](https://www.microsoft.com/en-us/translator/apps/features/)",
    "正式视频会议：如果客户用 Zoom，可用 translated captions；如果客户用 Teams，可用 live translated captions。Zoom 官方说明字幕翻译已支持 35 种语言；微软官方说明 Teams Premium 可为会议参与者提供 live translated captions。[Zoom translated captions](https://support.zoom.com/hc/en/article?id=zm_kb&sysparm_article=KB0059081) [Zoom blog](https://www.zoom.com/en/blog/translated-captions/) [Teams captions](https://learn.microsoft.com/en-us/microsoftteams/meeting-transcription-captions)",
    "中国团队侧会议留痕：腾讯会议 / VooV Meeting 或讯飞听见。腾讯会议帮助中心和官方新闻都写到，字幕、实时转写和会后录制页支持翻译，且国际版 VooV Meeting 已在 130+ 国家和地区上架；讯飞听见官方则强调中英文转写、中英文互译、会议记录成稿和私有云部署能力。[腾讯会议帮助中心](https://meeting.tencent.com/support/) [腾讯会议版本更新](https://meeting.tencent.com/support/version/) [VooV Meeting](https://meeting.tencent.com/news/zgqyyfch20240905.html) [讯飞听见智能会议系统](https://www.iflyrec.com/html/products/znhy.html)",
]

ORG = [
    "老板层：只看中文。所有英文来往都必须在内部沉淀成一页中文摘要，内容固定为：客户是谁、需求是什么、风险在哪里、建议怎么回、要不要接。",
    "邀约岗位：不要求自由英文表达，只做模板化动作。可以用英文脚本 + 翻译软件发首封邮件、发 LinkedIn 私信、发预约链接、收集客户资料；但一开始不要让他单独打英文 discovery call。",
    "会议层：老板如果必须参会，安排‘一个会说业务的中文负责人 + 一个会工具的人 + 实时字幕翻译’即可。老板负责判断，工具负责把英文变中文，不必让老板自己听英文。",
    "文书层：报价单、 engagement letter、官网文案、正式法律分析，不要一键机翻直接发。必须经过人审，哪怕是兼职双语人员或外部英文编辑审一次。",
    "团队层：最值得补的不是‘全员英语培训’，而是补 1 个双语项目经理/外部兼职英文顾问。你们不需要一开始请很多美国人，但至少要有一个人能兜底外部英文表达。",
]

PLAN = [
    "低成本启动版：Edge + DeepL + Outlook/Word 翻译 + Zoom/Teams 字幕翻译。适合先试市场、先接线索。",
    "标准执行版：再加一个兼职双语项目经理，负责审邮件、跟会、改文案、把英文客户需求翻回中文。这是我最推荐的平衡方案。",
    "稳妥扩张版：当美国业务有稳定线索后，再补专职双语 BD/客服或与美国端律师、顾问建立固定合作。老板仍然可以主要看中文。",
]

RULES = [
    "不要把机器翻译当法律结论。它适合邀约、初聊、会议理解、纪要整理，不适合未经复核就发正式法律意见。",
    "要建术语表。比如 mainland China、PRC law、arbitration、preservation、enforcement、notarization、apostille、supplier verification 这类词，一旦统一，团队英文稳定度会明显提升。DeepL 官方支持 glossary，适合做这一层。[DeepL glossary](https://support.deepl.com/hc/en-us/articles/360021634540-About-the-glossary)",
    "所有外部沟通都做中英双留痕。英文原文保留，内部自动生成中文摘要，避免老板和律师决策时信息失真。",
    "先从文字沟通做起，再进入英文会议。文字沟通更容易用模板和翻译工具控风险。",
]


def build_docx():
    doc = init_doc()
    add_title(doc, "美国市场语言执行补充方案", "适用于：老板不懂英文、邀约岗位英文较弱，但仍要启动美国客户业务")

    add_heading(doc, "一、核心结论")
    add_bullets(doc, SUMMARY)

    add_heading(doc, "二、推荐的软件组合")
    add_bullets(doc, STACK)

    add_heading(doc, "三、最稳的组织做法")
    add_bullets(doc, ORG)

    add_heading(doc, "四、建议采用的执行版本")
    add_bullets(doc, PLAN)

    add_heading(doc, "五、必须守住的底线")
    add_bullets(doc, RULES)

    add_heading(doc, "六、最后判断")
    add_para(doc, "老板不懂英文，并不是这个项目不能做的理由。真正决定能不能做成的，是你们有没有把‘英文接待、会议理解、文书输出、内部决策’拆成流程。先用工具把理解成本打下来，再用一个双语角色把关键节点兜住，这个项目就能启动。")
    add_para(doc, "建议：先按‘标准执行版’推进，也就是工具先上，同时尽快补一个兼职或外包的双语项目经理。这样投入不大，但成功率会明显高很多。")

    add_para(doc, "调研时间：2026 年 3 月 22 日；信息来源：各产品官方帮助中心、官方功能页与官方公告。", style="Small")
    doc.save(DOCX_OUT)


def build_md():
    parts = [
        "# 美国市场语言执行补充方案",
        "",
        "适用于：老板不懂英文、邀约岗位英文较弱，但仍要启动美国客户业务",
        "",
        "## 一、核心结论",
        "",
    ]
    for item in SUMMARY:
        parts.append(f"- {item}")
    parts += ["", "## 二、推荐的软件组合", ""]
    for item in STACK:
        parts.append(f"- {item}")
    parts += ["", "## 三、最稳的组织做法", ""]
    for item in ORG:
        parts.append(f"- {item}")
    parts += ["", "## 四、建议采用的执行版本", ""]
    for item in PLAN:
        parts.append(f"- {item}")
    parts += ["", "## 五、必须守住的底线", ""]
    for item in RULES:
        parts.append(f"- {item}")
    parts += [
        "",
        "## 六、最后判断",
        "",
        "老板不懂英文，并不是这个项目不能做的理由。真正决定能不能做成的，是你们有没有把‘英文接待、会议理解、文书输出、内部决策’拆成流程。先用工具把理解成本打下来，再用一个双语角色把关键节点兜住，这个项目就能启动。",
        "",
        "建议：先按‘标准执行版’推进，也就是工具先上，同时尽快补一个兼职或外包的双语项目经理。这样投入不大，但成功率会明显高很多。",
        "",
        "调研时间：2026 年 3 月 22 日；信息来源：各产品官方帮助中心、官方功能页与官方公告。",
        "",
    ]
    MD_OUT.write_text("\n".join(parts), encoding="utf-8")


if __name__ == "__main__":
    build_md()
    build_docx()
