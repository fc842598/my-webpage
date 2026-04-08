from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\1\Desktop\家里用的图标")
DOCX_OUT = Path(r"C:\Users\1\Desktop\美国市场最终执行方案_终稿_20260322.docx")
PDF_DIR = Path(r"C:\Users\1\Desktop")
MD_OUT = ROOT / "output" / "doc" / "美国市场最终执行方案_终稿_20260322.md"


def set_ea_font(run, name: str = "微软雅黑", size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False, size: float = 9.1) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_ea_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


SOURCES = [
    ("S1", "官方", "美国 Census：February 2026 story", "https://www.census.gov/about/history/stories/monthly/2026/february-2026.html", "2024 年美国 Chinese population 为 4,803,516。"),
    ("S2", "官方", "美国 Census：Language at home ACS", "https://www.census.gov/newsroom/press-releases/2023/language-at-home-acs-5-year.html", "2018-2022 中，Chinese speakers 中 48.2% speak English very well。"),
    ("S3", "官方", "美国 Census：Top Trading Partners 2024", "https://www.census.gov/foreign-trade/statistics/highlights/top/top2412yr.html", "2024 年美国对华货物贸易总额约 5825 亿美元。"),
    ("S4", "官方", "美国 Census：U.S. trade in goods with China", "https://www.census.gov/foreign-trade/balance/c5700.html", "Census 当前页面仍显示 2025 年美国对华货物贸易规模很高。"),
    ("S5", "官方", "深圳市涉外法律服务指引", "https://sf.sz.gov.cn/ztzl/swflfwzl/zxxx/content/post_12240394.html", "深圳官方持续建设涉外法律服务能力。"),
    ("S6", "官方", "深圳企业海外综合服务平台两大法律专区正式上线", "https://sf.sz.gov.cn/gkmlpt/content/12/12552/post_12552338.html", "深圳官方已把涉外法律服务与合规指引纳入企业海外综合服务平台。"),
    ("S7", "官方", "深圳召开全市涉外法治建设暨企业合规建设工作推进会", "https://sf.sz.gov.cn/gkmlpt/content/12/12094/post_12094226.html", "深圳 2025 年继续强调涉外法治建设与涉外法律服务效能。"),
    ("S8", "官方", "SCIA Bay-to-Bay cooperation", "https://en.scia.com.cn/2024-03/13/c_983171.htm", "SCIA 与硅谷仲裁力量协作，利于深圳做跨境商事争议背书。"),
    ("S9", "官方", "California FLC FAQ", "https://www.calbar.ca.gov/Admissions/Special-Admissions/Foreign-Legal-Consultants-FLC/FAQ", "外国法律顾问可就其本国法执业，但不得执业 California law。"),
    ("S10", "官方", "California FLC page", "https://www.calbar.ca.gov/admissions/special-admissions/foreign-legal-consultants", "在 California 设点前，可评估 FLC 路径。"),
    ("S11", "官方", "New York Part 521", "https://www.nycourts.gov/ctapps/521rules10.htm", "纽约允许 foreign legal consultant 开设办公室并受限执业。"),
    ("S12", "官方", "Texas FLC certification", "https://www.texasbar.com/AM/Template.cfm?ContentID=48647&Template=%2FCM%2FHTMLDisplay.cfm", "Texas 同样存在 FLC certification 机制。"),
    ("S13", "官方", "Texas FLC application", "https://ble.texas.gov/m/foreign-legal-consultant-application", "Texas 要求 good standing 和 discipline statement。"),
    ("S14", "官方", "HCCH Apostille authority for China", "https://www.hcch.net/en/states/authorities/details3/?aid=1193", "中国内地 Apostille 已可用，且深圳外办在授权列表中。"),
    ("S15", "官方", "Microsoft Edge 翻译", "https://support.microsoft.com/en-us/topic/use-microsoft-translator-in-microsoft-edge-browser-4ad1c6cb-01a4-4227-be9d-a81e127fcb0b", "支持网页自动识别与翻译。"),
    ("S16", "官方", "Outlook 翻译", "https://support.microsoft.com/en-us/office/translator-in-outlook-for-windows-3d7e12ed-99d6-406e-a453-b9db0d9653fa", "支持邮件整封翻译和自动翻译。"),
    ("S17", "官方", "DeepL for Microsoft Word", "https://support.deepl.com/hc/en-us/articles/9851524013468-DeepL-for-Microsoft-Word", "支持 Word / Outlook / PowerPoint 插件翻译与润色。"),
    ("S18", "官方", "DeepL glossary", "https://support.deepl.com/hc/en-us/articles/360021634540-About-the-glossary", "支持术语表统一。"),
    ("S19", "官方", "Zoom translated captions", "https://www.zoom.com/en/blog/translated-captions/", "Zoom translated captions 支持 35 种语言。"),
    ("S20", "官方", "Teams transcription and captions", "https://learn.microsoft.com/en-us/microsoftteams/meeting-transcription-captions", "Teams 支持实时转录与 live translated transcription。"),
    ("S21", "官方", "腾讯会议出海能力", "https://meeting.tencent.com/news/zgqyyfch20240905.html", "VooV Meeting 已在 130+ 国家和地区上架。"),
    ("S22", "竞品", "Harris Sliwoski China law", "https://harris-sliwoski.com/practice-areas/china-law/", "国际化品牌所样本。"),
    ("S23", "竞品", "ChinaLawyers", "https://chinalawyers.net/", "英文直客页型强竞争者。"),
    ("S24", "竞品", "LawFirmChina", "https://www.lawfirmchina.com/", "Supplier validation 和 free consultation 导向明显。"),
    ("S25", "竞品", "CLN Debt engagement letter", "https://www.clndebt.com/images/LAsubK.pdf", "公开 flat fee 及尽调/催收打法。"),
]


EXEC_SUMMARY = [
    "最终判断：这个生意可做，建议立项执行；但执行路线必须从“泛美国、纯英文、全业务”改成“分阶段、双语、以中国内地法服务为核心”的版本。",
    "反查后的结论不是“市场不存在”，而是“市场存在，但最容易踩坑的地方在语言和合规边界”。如果不改打法，项目容易死在英语接待、页面转化和范围失控；如果按修正版路线走，成功率会明显上升。",
    "建议采用的终版路线是：第一阶段先做‘在美国的中文/双语客户 + 需要中国内地法支持的美国专业转介人’，第二阶段再扩张到英文直客页和英语直接获客，第三阶段再评估美国办公室或 FLC 路径。",
    "这意味着老板和邀约岗位不需要立刻变成英文高手。更稳的做法，是用翻译工具、术语表、英文模板、实时字幕和一个双语项目经理，把英语变成组织能力而不是个人天赋。",
]


REVERSE_CHECK = [
    ["市场需求", "8.5/10", "可行", "中国相关商事与文书需求长期存在；美国华裔与中文使用者体量足够，且中美贸易仍大 [S1][S2][S3][S4]。"],
    ["竞争环境", "7.5/10", "可行但要聚焦", "英文搜索竞争真实存在，但主要是 landing page 型团队在吃流量，不是所有大所都在抢直客 [S22][S23][S24][S25]。"],
    ["合规边界", "6.5/10", "可控", "只要明确提供的是 PRC law / China mainland legal support，不碰 U.S. law，风险可控；未来若设美国办公室或常驻人员，应先做 FLC/当地合规评估 [S9]-[S13]。"],
    ["语言与交付", "4.5/10（现状）", "现状不足", "若继续走纯英文自由发挥模式，你们当前团队英语能力不够稳。"],
    ["语言与交付", "7.8/10（修正后）", "可行", "采用双语模板、实时翻译和双语项目经理后，邀约、会议和文书都能跑起来 [S15]-[S21]。"],
]


WHY_FEASIBLE = [
    "美国市场里并不只有纯英语客户。Census 2024 数据显示，美国 Chinese population 已达到 4,803,516 [S1]；而 Census 2023 口径显示，在 2018-2022 ACS 中，Chinese speakers 只有 48.2% 表示自己 English “very well” [S2]。这意味着，‘美国市场’里面本身就存在一个对中文或双语服务更友好的切片。",
    "中美之间与中国内地相关的法律需求并未消失。美国 Census 公开数据表明，2024 年美国对华货物贸易总额仍约 5825 亿美元 [S3]；Census 当前页面显示 2025 年规模仍高 [S4]。只要贸易、采购、供应链、付款和跨境文件仍在发生，中国内地法支持就仍有需求。",
    "竞争者已经证明需求可以被转化。Harris Sliwoski、ChinaLawyers、LawFirmChina、CLN Debt 都在用英文或双语页面承接‘中国供应商纠纷、合同、追债、尽调’这类高意向需求 [S22]-[S25]。市场不是空白，但也不是没有入口。",
    "深圳有供给优势。深圳官方持续建设涉外法律服务体系，SCIA 与硅谷仲裁合作也为深圳输出国际商事争议能力提供了背书 [S5]-[S8]。这类供给条件本身就是进入美国市场的基础设施。"
]


WHAT_TO_CHANGE = [
    ["原来思路", "修正版终稿", "为什么必须改"],
    ["一上来做英语美国市场", "先做在美国的中文/双语客户，再逐步扩英语直客", "语言成本更低，起盘更稳，询盘更容易拿到。"],
    ["一开始什么业务都接", "先打中国供应商/合同/货款/仲裁四类商事需求", "更高频、更高客单价、更适合深圳供给。"],
    ["老板和邀约岗直接顶英文", "老板只看中文决策，邀约岗只跑模板和预约，双语 PM 兜底", "避免把项目死在语言环节。"],
    ["只做英文官网等客户来", "官网 + landing pages + Search test + referral 一起上", "法律业务单靠官网通常不够。"],
    ["过早考虑美国办公室", "先远程交付验证，后评估 California / New York / Texas 的 FLC 路径", "先证明业务，再决定是否重投入。"],
]


TARGET_SEGMENTS = [
    ["阶段 1 主客群", "在美国的中文/双语企业主、进口商、跨境卖家、华人创业者", "中国供应商纠纷、采购合同、货款追收、主体核验、仲裁/诉讼协调"],
    ["阶段 1 主客群", "美国律师、CPA、贸易顾问、货代、采购顾问", "把你们当 China counsel / PRC law support 使用"],
    ["阶段 2 客群", "英语为主的美国 SME 与品牌方", "当 landing pages、英文交付、案例积累成熟后再放大"],
    ["暂不优先", "婚姻家事、继承、房产、私人客户全覆盖", "可以做，但不宜当作第一阶段主战场"],
]


PRODUCTS = [
    ["48 小时中国内地法初判", "快速判断能不能追、怎么追、先做什么", "¥999 - ¥2,999"],
    ["供应商核验与主体调查", "工商、裁判、地址、联系人、关联主体、执行风险", "¥6,000 - ¥15,000"],
    ["双语 NNN / 采购合同包", "交易前预防争议", "¥8,000 - ¥25,000"],
    ["商账催收 / 索赔包", "货款、预付款、赔偿追收", "¥10,000 - ¥30,000 + success fee"],
    ["中国仲裁/诉讼项目管理包", "真正进入程序时的落地服务", "¥50,000 起"],
    ["中国法律月顾问", "长期在华经营、采购、合规支持", "¥20,000 - ¥50,000 / 月"],
]


CHANNELS = [
    ["中文 landing pages（美国场景）", "第 1 阶段主力", "降低语言门槛，优先打在美中文/双语客户"],
    ["英文 landing pages", "第 1 阶段同步但不做过多", "先做 2 - 3 个高意向页，不追求全站铺满"],
    ["Google Search 小规模试投", "第 1 阶段验证", "只投 supplier dispute / debt recovery / China contract / Chinese supplier verification"],
    ["Referral network", "必须同步", "找美国律师、CPA、贸易顾问、采购顾问、华商组织合作"],
    ["WeChat / WhatsApp / 邮件预约", "第 1 阶段标配", "联系方式必须低门槛，不要只留邮箱表单"],
    ["内容营销", "第 2 阶段放大", "围绕真实问题写 FAQ、案例、证据清单、流程说明"],
]


LANGUAGE_MODEL = [
    "老板：不直接处理英文原文。所有对外英文沟通都先沉淀为中文摘要，摘要固定字段为‘客户是谁、现在卡在哪、建议怎么回、是否值得接’。",
    "邀约岗位：不要求自由英文表达。只负责用双语模板发首封邮件、发邀约信息、发预约链接、收集资料、推动客户进表单和进会议。",
    "双语项目经理：这是终稿里的新增关键角色。可以是兼职、外包或顾问，但必须有人负责英文邮件校对、会前准备、会议跟进、会后纪要和报价语气统一。",
    "律师：重点负责 PRC law 实体判断与交付，不要被迫兼做所有英文接待与商务跟单。",
    "外部 U.S. counsel / partner counsel：只在需要美国法或美国本地程序意见时介入，既能控风险，也能作为转介绍网络的一部分。"
]


TOOLS = [
    ["Microsoft Edge 翻译", "看英文网站、竞品页、法院/政府信息", "内置翻译，启动快 [S15]"],
    ["Outlook Translator", "收英文邮件、回英文邮件前理解原文", "支持整封邮件翻译和自动翻译 [S16]"],
    ["DeepL for Word / Outlook", "写英文文书、润色英文对外文本", "Word / Outlook 插件 + 术语表 [S17][S18]"],
    ["Zoom translated captions", "和美国客户正式视频会议", "35 种语言实时字幕翻译 [S19]"],
    ["Microsoft Teams live transcription / translated transcription", "客户如果偏好 Teams", "会议内实时转录与翻译 [S20]"],
    ["腾讯会议 / VooV", "中国团队内部协同和出海会议备用", "国际版上架 130+ 国家和地区 [S21]"],
]


COMPLIANCE_RULES = [
    "终稿判断：第一阶段可以先不在美国设办公室，也不让美国客户误以为你们在提供 U.S. law。对外统一表述必须是：`China Mainland Legal Support`、`PRC law advice`、`China dispute / contract / enforcement support`。",
    "California、New York、Texas 都存在 Foreign Legal Consultant 路径 [S9]-[S13]。这说明美国监管体系承认“外国法意见”与“州法/联邦法意见”之间的边界。基于这些规则，我的推断是：第一阶段你们最可行的路径不是急着落美国办公室，而是先把服务范围严格限制在中国内地法。",
    "所有官网、报价单、 engagement letter 和邮件签名都要写清楚：不提供 U.S. law advice；如涉及 U.S. law matters，将由本地持证律师另行提供或联合提供。这一条必须写进模板。",
    "如果未来要在 California / New York / Texas 长期驻点、开办公室、驻场 BD 或长期雇员对外承接当地法律业务，则必须在落地前再做当地律师合规复核。"
]


NINETY_DAY = [
    ["第 1 - 7 天", "定服务边界、统一对外口径、设术语表、确定价格带、补双语项目经理", "没有这一层，不要上线英文获客"],
    ["第 8 - 14 天", "做 2 个中文美国场景页 + 2 个英文高意向页，配置表单、邮箱、预约、WhatsApp/微信入口", "先做能转化的，不先做大而全官网"],
    ["第 15 - 30 天", "完成邀约模板、报价模板、会前问卷、证据清单；开始触达首批 referral partners", "邀约岗位按模板执行"],
    ["第 31 - 45 天", "小规模 Search test 上线；开始接 discovery calls；双语 PM 跟会并沉淀纪要", "重点看有效询盘而不是点击量"],
    ["第 46 - 60 天", "优化页面和广告词；补案例、FAQ、流程页；做付费咨询产品", "把免费咨询逐步升级成筛选型咨询"],
    ["第 61 - 90 天", "复盘线索质量、签约率、回款与团队承接能力；决定是否加预算和扩英文直客", "90 天后再决定是否放大"],
]


BUDGET = [
    ["轻量验证版", "¥100,000 - ¥160,000", "中文/英文核心页面、双语 PM 外包、少量投放、基础工具"],
    ["标准执行版", "¥180,000 - ¥320,000", "更完整页面、Search test、内容、双语 PM、Referral 拓展"],
    ["不建议首期支出", "美国办公室、重品牌片、大规模社媒铺投", "在验证前做这些，性价比很差"],
]


KPI = [
    ["30 天", "至少拿到 5 - 10 个有效询盘、2 - 3 个正式会议", "如果 30 天没有有效询盘，优先查页面与渠道，不先怪市场"],
    ["60 天", "至少拿到 10 - 20 个有效询盘、3 - 5 个付费咨询、1 - 2 个正式委托", "说明渠道开始成型"],
    ["90 天", "15 - 30 个有效询盘、5 - 10 个付费咨询、2 - 5 个正式委托", "达到这个区间，就值得扩大投入"],
]


IMMEDIATE_ACTIONS = [
    "先批准这份终稿，不再按‘泛美国英语市场’执行，而按‘分阶段、双语、以 PRC law 为核心’执行。",
    "本周内补 1 个双语项目经理角色，可以先外包，不必等全职到位。",
    "本周内把对外口径统一成 `China Mainland Legal Support`，并写进网站、邮件、报价单、 engagement letter。",
    "本周内先做 4 个页面：中文 2 个、英文 2 个。先上 supplier dispute、debt recovery、supplier verification、NNN / contract。",
    "邀约岗位本周开始训练的不是口语，而是模板执行、资料收集、预约推进和 CRM 录入。",
]


APPENDIX_NOTES = [
    "本文件是商业执行方案，不构成美国各州执业法律意见。",
    "涉及 California / New York / Texas FLC 的内容，是基于各州官方公开规则做的业务可行性判断；若未来在当地设办公室或常驻人员，应再让当地律师复核。",
    "关于广告情况，本次仍以公开搜索结果、落地页结构、CTA 设计和竞品公开信息为依据，不包含 Google Ads 后台 spend 数据。",
]


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    if "Small" not in doc.styles:
        style = doc.styles.add_style("Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(9)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("美国市场最终执行方案\n终稿")
    set_ea_font(r, size=22, bold=True, color=(31, 78, 121))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("适用对象：深圳律所面向在美客户提供中国内地法律支持")
    set_ea_font(r2, size=12.5)

    meta = [
        "版本：终稿 v1.0",
        "时间：2026 年 3 月 22 日",
        "用途：老板拍板后即可执行的工作文件",
        "核心结论：建议立项，但按修正版路径执行，不按泛英语路线硬冲",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(line)
        set_ea_font(r, size=10.5)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    if level == 1:
        set_ea_font(r, size=15, bold=True, color=(31, 78, 121))
    elif level == 2:
        set_ea_font(r, size=12.5, bold=True, color=(31, 78, 121))
    else:
        set_ea_font(r, size=11.2, bold=True, color=(31, 78, 121))


def add_para(doc: Document, text: str, style: str = "Normal", bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_ea_font(r1, size=9 if style == "Small" else 10.5, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_ea_font(r2, size=9 if style == "Small" else 10.5)
    else:
        r = p.add_run(text)
        set_ea_font(r, size=9 if style == "Small" else 10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.35
        rb = p.add_run("• ")
        set_ea_font(rb)
        rr = p.add_run(item)
        set_ea_font(rr)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], size: float = 9.0) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, size=size)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, size=size)
    doc.add_paragraph("")


def build_docx() -> None:
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)

    add_heading(doc, "一、最终结论")
    add_bullets(doc, EXEC_SUMMARY)

    add_heading(doc, "二、反查后的可行性结论")
    add_para(doc, "我对这个项目重新按‘市场、合规、执行’三条线做了反查。结论不是否定项目，而是把它从高风险版本修正为可执行版本。", bold_prefix="我对这个项目重新按‘市场、合规、执行’三条线做了反查。")
    add_table(doc, ["维度", "评分", "结论", "反查结果"], REVERSE_CHECK, size=8.8)

    add_heading(doc, "三、为什么这个项目仍然值得做")
    add_bullets(doc, WHY_FEASIBLE)

    add_heading(doc, "四、终稿相较前版的关键修改")
    add_table(doc, WHAT_TO_CHANGE[0], WHAT_TO_CHANGE[1:], size=8.9)

    add_heading(doc, "五、终稿推荐的市场切入顺序")
    add_para(doc, "终稿建议不再把‘美国市场’理解为一开始就做纯英语全国市场，而是拆成三层。这样既能降低语言门槛，也更符合你们目前的组织能力。")
    add_table(doc, ["客群层级", "客户画像", "核心需求"], TARGET_SEGMENTS, size=8.9)
    add_bullets(
        doc,
        [
            "第一阶段先做在美国的中文/双语客户，不代表只做华人，而是代表先打更容易转化、语言阻力更低的那部分客户。",
            "第二阶段再增加英语页面和英语直客，并把 referral network 做厚。",
            "第三阶段等询盘、签约、交付稳定后，再评估是否需要在 California / New York / Texas 走 FLC 或其他本地化路径。",
        ],
    )

    add_heading(doc, "六、推荐产品线")
    add_table(doc, ["产品", "解决什么问题", "价格带"], PRODUCTS, size=9.0)

    add_heading(doc, "七、获客渠道终稿")
    add_table(doc, ["渠道", "优先级", "为什么要做"], CHANNELS, size=8.9)
    add_para(doc, "终稿建议：不要把资源压在单一渠道上。第一阶段至少要同时跑‘页面 + 少量 Search test + referral’三条线。")

    add_heading(doc, "八、语言与组织执行模型")
    add_bullets(doc, LANGUAGE_MODEL)
    add_table(doc, ["工具", "主要用途", "为什么选它"], TOOLS, size=8.8)

    add_heading(doc, "九、合规边界与底线")
    add_bullets(doc, COMPLIANCE_RULES)

    add_heading(doc, "十、90 天执行计划")
    add_table(doc, ["时间段", "关键动作", "执行要求"], NINETY_DAY, size=8.8)

    add_heading(doc, "十一、预算与结果判断")
    add_table(doc, ["预算版本", "建议区间", "说明"], BUDGET, size=8.9)
    add_table(doc, ["检查节点", "目标", "如何解读"], KPI, size=8.9)

    add_heading(doc, "十二、老板现在需要拍板的事项")
    add_bullets(doc, IMMEDIATE_ACTIONS)

    add_heading(doc, "十三、最终建议")
    add_para(doc, "最终建议：批准立项，但只批准‘修正版终稿’。不要批准纯英语广撒网版本，也不要批准一上来做美国办公室或全业务出海版。")
    add_para(doc, "真正最稳的路径是：先用中文/双语市场切片在美国拿到第一批客户和转介绍，再逐步把英文交付、英文页面和美国端网络做起来。这样既不违背做美国市场的目标，也不会把团队一次性推到英语和合规的风险区。")
    add_para(doc, "一句话判断：这个项目可以做，而且值得做；但必须按这份终稿的节奏和边界执行。", bold_prefix="一句话判断：")

    add_heading(doc, "附录一：执行时必须记住的三条红线")
    add_bullets(
        doc,
        [
            "不要把中国内地法服务说成美国法服务。",
            "不要让不会英文的人单独自由发挥对外关键沟通。",
            "不要在 90 天验证前重投入办公室、品牌片和泛社媒投流。",
        ],
    )

    add_heading(doc, "附录二：主要公开来源")
    add_para(doc, "以下来源用于市场、合规与工具可行性反查。", style="Small")
    for code, category, title, url, note in SOURCES:
        add_para(doc, f"{code} [{category}] {title} - {url} - {note}", style="Small")

    add_heading(doc, "附录三：说明")
    add_bullets(doc, APPENDIX_NOTES)

    doc.save(DOCX_OUT)


def build_md() -> None:
    lines: list[str] = []
    lines.append("# 美国市场最终执行方案（终稿）")
    lines.append("")
    lines.append("适用对象：深圳律所面向在美客户提供中国内地法律支持")
    lines.append("")
    lines.append("## 一、最终结论")
    lines.append("")
    for item in EXEC_SUMMARY:
        lines.append(f"- {item}")
    lines.append("")

    lines.append("## 二、反查后的可行性结论")
    lines.append("")
    lines.append("| 维度 | 评分 | 结论 | 反查结果 |")
    lines.append("| --- | --- | --- | --- |")
    for row in REVERSE_CHECK:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 三、为什么这个项目仍然值得做")
    lines.append("")
    for item in WHY_FEASIBLE:
        lines.append(f"- {item}")
    lines.append("")

    lines.append("## 四、终稿相较前版的关键修改")
    lines.append("")
    lines.append("| 原来思路 | 修正版终稿 | 为什么必须改 |")
    lines.append("| --- | --- | --- |")
    for row in WHAT_TO_CHANGE[1:]:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 五、终稿推荐的市场切入顺序")
    lines.append("")
    lines.append("| 客群层级 | 客户画像 | 核心需求 |")
    lines.append("| --- | --- | --- |")
    for row in TARGET_SEGMENTS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 六、推荐产品线")
    lines.append("")
    lines.append("| 产品 | 解决什么问题 | 价格带 |")
    lines.append("| --- | --- | --- |")
    for row in PRODUCTS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 七、获客渠道终稿")
    lines.append("")
    lines.append("| 渠道 | 优先级 | 为什么要做 |")
    lines.append("| --- | --- | --- |")
    for row in CHANNELS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 八、语言与组织执行模型")
    lines.append("")
    for item in LANGUAGE_MODEL:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("| 工具 | 主要用途 | 为什么选它 |")
    lines.append("| --- | --- | --- |")
    for row in TOOLS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 九、合规边界与底线")
    lines.append("")
    for item in COMPLIANCE_RULES:
        lines.append(f"- {item}")
    lines.append("")

    lines.append("## 十、90 天执行计划")
    lines.append("")
    lines.append("| 时间段 | 关键动作 | 执行要求 |")
    lines.append("| --- | --- | --- |")
    for row in NINETY_DAY:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 十一、预算与结果判断")
    lines.append("")
    lines.append("| 预算版本 | 建议区间 | 说明 |")
    lines.append("| --- | --- | --- |")
    for row in BUDGET:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("| 检查节点 | 目标 | 如何解读 |")
    lines.append("| --- | --- | --- |")
    for row in KPI:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")

    lines.append("## 十二、老板现在需要拍板的事项")
    lines.append("")
    for item in IMMEDIATE_ACTIONS:
        lines.append(f"- {item}")
    lines.append("")

    lines.append("## 十三、最终建议")
    lines.append("")
    lines.append("最终建议：批准立项，但只批准“修正版终稿”。不要批准纯英语广撒网版本，也不要批准一上来做美国办公室或全业务出海版。")
    lines.append("")
    lines.append("真正最稳的路径是：先用中文/双语市场切片在美国拿到第一批客户和转介绍，再逐步把英文交付、英文页面和美国端网络做起来。")
    lines.append("")
    lines.append("一句话判断：这个项目可以做，而且值得做；但必须按这份终稿的节奏和边界执行。")
    lines.append("")

    lines.append("## 附录：主要公开来源")
    lines.append("")
    for code, category, title, url, note in SOURCES:
        lines.append(f"- {code} [{category}] {title} - {url} - {note}")
    lines.append("")
    for note in APPENDIX_NOTES:
        lines.append(f"- {note}")
    lines.append("")

    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    build_md()
    build_docx()


if __name__ == "__main__":
    main()
